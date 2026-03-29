import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc = docx.Document()
doc.add_heading('Financial Bubble Detection System: Project Report', 0)

doc.add_heading('1. Project Overview & What We Did', level=1)
doc.add_paragraph("We built a near real-time financial market intelligence dashboard and machine learning classification system focused on the Indian stock market (specifically NIFTY 50 and NIFTY 500). The system detects whether the current market state is in a Normal, Bubble, or Crash phase.")
doc.add_paragraph("The project successfully integrates:")
doc.add_paragraph("• Real-time Price Engine: Fetches live data and calculates deep technical indicators (e.g., MACD, RSI, Volatility, Drawdowns).")
doc.add_paragraph("• Macroeconomic Engine: Tracks longer-term macro features such as GDP growth, CPI inflation, and Repo rates.")
doc.add_paragraph("• Sentiment Engine: Fetches real-time financial news and passes them through an NLP model to gauge market sentiment.")
doc.add_paragraph("• Ensemble Machine Learning AI: Aggregates all these diverse feature sets to predict the probability of a market crash or bubble.")

doc.add_heading('2. Machine Learning Models & Training Methodology', level=1)
doc.add_paragraph("The system utilizes robust tree-based classification models designed to handle tabular financial data.")
doc.add_heading('Models Used', level=2)
doc.add_paragraph("1. XGBoost (Extreme Gradient Boosting): Leverages gradient boosted decision trees for high accuracy and fast inference.")
doc.add_paragraph("2. Random Forest Classifier: Uses an ensemble of deep decision trees to ensure stability and reduce overfitting.")
doc.add_paragraph("3. Stacking Ensemble: Combines the probabilities from both XGBoost and Random Forest into an overarching \"meta-model\" to capture the strengths of both algorithms and make the final decision.")

doc.add_heading('Training Methodology', level=2)
doc.add_paragraph("Financial crash and bubble events are extremely rare compared to normal market days, creating a severe class imbalance problem.")
doc.add_paragraph("1. Feature Engineering: Calculates extensive regime features (log returns, PSY (Psychological Line), 10-day & 63-day realized volatility, volatility z-scores, 21-day return skews/kurtosis, MACD, RSI).")
doc.add_paragraph("2. Oversampling with ADASYN: During training, the ADASYN (Adaptive Synthetic) algorithm is applied exclusively to the training set. It generates synthetic data points for the minority \"Crash\" and \"Bubble\" classes, upsampling them so the models do not ignore them and learn their distinct patterns.")
doc.add_paragraph("3. Walk-Forward Validation: Because financial data is strictly time-series dependent, standard cross-validation does not work. The models are validated using varying rolling time windows (e.g., train on 2008–2019, test on 2019-2020, then step forward to the next years).")
doc.add_paragraph("4. Target Metric: The models are optimized for the Macro F1-Score to ensure that all classes (Crash, Bubble, Normal) perform equally well without letting the majority \"Normal\" class artificially inflate the accuracy.")

doc.add_heading('3. Detecting Bubbles Based on Scores', level=1)
doc.add_paragraph("The system defines bubbles and crashes organically using advanced statistical scoring techniques before the ML models learn to predict them. This core labeling logic is strictly mathematical and is found in zscore_labeling.py:")
doc.add_paragraph("1. Z-Score Calculation: The primary foundation is the Z-Score of the closing price using a 30-day rolling window (representing standard deviations from the mean).")
doc.add_paragraph("    • Crash Threshold: Automatically labeled as a \"Crash\" if the Z-Score drops extremely fast (below -2.0).")
doc.add_paragraph("    • Bubble Threshold: Automatically labeled as a \"Bubble\" if the Z-Score spikes (above +2.0).")
doc.add_paragraph("2. BSADF (Backward Sup-ADF) Scoring: For a more mathematically rigorous bubble detection, the system computes the Backward Sup-Augmented Dickey-Fuller (BSADF) statistic on the log-prices.")
doc.add_paragraph("    • The BSADF test is an econometric approach designed by economists to explicitly detect \"explosive\" (exponential) growth behavior typical in financial bubbles.")
doc.add_paragraph("    • If the current BSADF score crosses the 95th percentile upper-bound (calculated over a 504-day rolling window) AND the market return is positive, the day is rigorously labeled as a \"Bubble\" even if the Z-Score is moderate.")
doc.add_paragraph("The Machine Learning models are trained on these historical labels. In live inference, the ensemble model outputs the Probability (0% to 100%) for each category. A final strict safety filter is applied in the app: if the Crash/Bubble probability is extremely high (e.g., >97% for recent windows) and comfortably beats out the second-highest probability, the live alert is triggered.")

doc.add_heading('4. Sentiment Index Calculation & Integration', level=1)
doc.add_paragraph("Market sentiment is a powerful leading indicator of irrational exuberance (bubbles) or mass panic (crashes).")
doc.add_heading('How it is calculated', level=2)
doc.add_paragraph("1. News Fetching: The news_fetcher.py script retrieves recent financial news headlines relating to the NIFTY 50.")
doc.add_paragraph("2. FinBERT NLP Processing: The headlines are passed into ProsusAI/finbert, a deep learning NLP model fine-tuned entirely on financial communication data.")
doc.add_paragraph("3. Polarity Scoring: FinBERT outputs a label (positive, negative, or neutral) and a confidence score for each headline.")
doc.add_paragraph("    • Positive headlines receive a value equal to their confidence score (e.g., +0.85).")
doc.add_paragraph("    • Negative headlines receive a negative value (e.g., -0.90).")
doc.add_paragraph("    • Neutral headlines receive a score of 0.0.")
doc.add_paragraph("4. Daily Sentiment Index: The scores of all news articles in a single day are averaged together to create a single continuous variable known as the daily_sentiment_index. A 3-day rolling mean (sentiment_momentum) is also computed to track smooth trends.")
doc.add_heading('How it is used in calculating a Bubble/Crash', level=2)
doc.add_paragraph("The daily_sentiment_index is injected directly into the feature matrix alongside technical and macroeconomic indicators.")
doc.add_paragraph("• During training, the models inherently learn the historical relationships (e.g., high euphoria/ultra-positive sentiment usually aligns with the peak of a Bubble; collapsing sentiment accelerates a Crash).")
doc.add_paragraph("• During live inference (app.py), the model fuses the daily sentiment index with the technical MACD/Z-scores and macro rates to produce the final confidence vote for Crash, Bubble, or Normal predictions.")
doc.add_paragraph("• Fallback Heuristic Plan: If the ML model fails to load, the raw negative polarity of the sentiment index is programmed to linearly and directly increase the \"Crash Risk\" probability of the market.")

doc.save('/Users/sohinsanthosh/bubble_detection_miniProject/mini_project/report_bubble_detection.docx')
print("Saved to /Users/sohinsanthosh/bubble_detection_miniProject/mini_project/report_bubble_detection.docx")
